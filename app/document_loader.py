from pypdf import PdfReader
from typing import Literal
from pathlib import Path
import docx


DocType = Literal["documentation", "rfc", "research", "manual"]

class DocumentLoader:
    
    FOLDER_TYPE_MAPPING = {
        "documentations": "documentation",
        "documentation": "documentation",
        "rfc": "rfc",
        "rfcs": "rfc",
        "research_papers": "research",
        "research": "research",
        "papers": "research",
        "software_manuals": "manual",
        "manuals": "manual",
        "manual": "manual"
    }

    def __init__(self, file_path: str) -> None:
        self.file_path = Path(file_path)
        self.doc_type = self._extract_doc_type()
        self.source =self._extract_source()

    def _extract_doc_type(self) -> DocType:
        
        parts = [p.lower() for p in self.file_path.parts]
        for part in parts:
            if part in self.FOLDER_TYPE_MAPPING:
                return self.FOLDER_TYPE_MAPPING[part]
        
        return "documentation"
    
    def _extract_source(self) -> str:
        parts = list(self.file_path.parts)

        if self.doc_type == "rfc":
            import re
            match = re.search(r'rfc\s*(\d+)', self.file_path.name.lower())
            if match:
                return f"rfc{match.group(1)}"
        
        doc_type_folder = None
        for i, part in enumerate(parts):
            if part.lower() in self.FOLDER_TYPE_MAPPING:
                doc_type_folder = i
                break
        
        #if in case there's a subfolder after doc_type folder, use that
        if doc_type_folder is not None and doc_type_folder + 1 < len(parts) - 1:
            return parts[doc_type_folder + 1].lower()
        
        return self.file_path.stem.lower()
    def load_pdf(self) -> str:
        try:
            reader = PdfReader(str(self.file_path))
            text = []
            for page in reader.pages:
                text.append(page.extract_text())
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Error reading PDF: {e}")
    
    def load_txt(self, encoding="utf-8") -> str:
        with open(str(self.file_path), "r", encoding=encoding) as f:
            return f.read()
    
    def load_docx(self) -> str:
        try:
            doc = docx.Document(str(self.file_path))
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    fullText.append(row_text)
            
            return '\n'.join(fullText)
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {e}")

    def load_md(self) -> str:
        with open(self.file_path, "r") as f:
            return f.read()
    
    def load(self) -> str:
        
        if not self.file_path:
            raise ValueError("File path is required")

        extension = self.file_path.suffix.lstrip(".").lower()

        loaders = {
            "pdf": self.load_pdf,
            "txt": self.load_txt,
            "docx": self.load_docx,
            "md": self.load_md
        }

        loader = loaders.get(extension)
        if not loader:
            raise ValueError(f"Unsupported file format: {self.file_path}")
        
        return loader()
    
    def get_metadata(self) -> dict:
        return {
            "doc_type": self.doc_type,
            "source": self.source,
            "file_path": str(self.file_path),
            "file_name": self.file_path.name
        }